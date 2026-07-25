from io import BytesIO
from pathlib import Path
import shutil

from docx import Document
from pypdf import PdfReader

from recruitment.config import get_settings

SUPPORTED_SUFFIXES = {".pdf", ".docx", ".txt"}


def extract_text_from_file(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        text = "\n\n".join(page.extract_text() or "" for page in PdfReader(str(path)).pages)
        return text.strip() or _ocr_pdf(path)
    if suffix == ".docx":
        document = Document(str(path))
        text = _docx_text(document)
        return text.strip() or _ocr_docx_images(document)
    if suffix == ".txt":
        return path.read_text(encoding="utf-8", errors="ignore")
    raise ValueError(f"Unsupported file type: {suffix or '(none)'}")


def _docx_text(document: Document) -> str:
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    table_cells = [
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    ]
    return "\n".join(paragraphs + table_cells)


def _ocr_pdf(path: Path) -> str:
    try:
        import fitz
    except ImportError as error:
        raise ValueError("OCR support requires PyMuPDF and pytesseract") from error
    document = fitz.open(path)
    try:
        images = []
        for page in document:
            pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
            images.append(pixmap.tobytes("png"))
        return _ocr_images(images)
    finally:
        document.close()


def _ocr_docx_images(document: Document) -> str:
    images = [
        relationship.target_part.blob
        for relationship in document.part.rels.values()
        if "image" in relationship.reltype
    ]
    return _ocr_images(images)


def _ocr_images(images: list[bytes]) -> str:
    if not images:
        return ""
    try:
        import pytesseract
        from PIL import Image
    except ImportError as error:
        raise ValueError("OCR support requires PyMuPDF and pytesseract") from error

    command = _tesseract_command()
    if command is None:
        raise ValueError("Tesseract OCR is not installed or configured")
    pytesseract.pytesseract.tesseract_cmd = str(command)
    try:
        return "\n\n".join(
            pytesseract.image_to_string(
                Image.open(BytesIO(image)), lang=get_settings().tesseract_languages
            )
            for image in images
        ).strip()
    except pytesseract.TesseractNotFoundError as error:
        raise ValueError("Tesseract OCR is not installed or configured") from error
    except pytesseract.TesseractError as error:
        raise ValueError(f"Tesseract OCR failed: {error}") from error


def _tesseract_command() -> Path | None:
    settings = get_settings()
    candidates = [
        settings.tesseract_cmd,
        Path(shutil.which("tesseract")) if shutil.which("tesseract") else None,
        Path(r"C:\Program Files\Tesseract-OCR\tesseract.exe"),
    ]
    return next((candidate for candidate in candidates if candidate and candidate.exists()), None)
